import os
import pandas as pd
import re
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path

def process_excel_report():
    # Get tomorrow's date in various formats
    tomorrow = datetime.now() + timedelta(days=1)
    tomorrow_date_str = tomorrow.strftime("%m%d%Y")
    tomorrow_formatted = tomorrow.strftime("%A, %B %d, %Y")
    tomorrow_file_date = tomorrow.strftime("%Y%m%d")
    
    # Define file paths
    downloads_folder = r"C:\Users\user1\Downloads\WebScrapedData"
    
    # Find the Excel file with tomorrow's date
    expected_filename = f"MART_Trips_{tomorrow_file_date}.xlsx"
    input_file = os.path.join(downloads_folder, expected_filename)
    
    if not os.path.exists(input_file):
        print(f"Excel file {expected_filename} not found in {downloads_folder}")
        return
    
    output_file = os.path.join(downloads_folder, f"Schedule_report_{tomorrow_date_str}.docx")
    
    print(f"Processing file: {input_file}")
    
    # Read the Excel file
    df = pd.read_excel(input_file)
    
    # Step 1: Rename headers
    headers = [
        "Trip Accepted", "Vehicle Type", "Name", "Phone", "Date", 
        "Trip Direction", "P/U Time", "Appt Time", "P/U Address/Entrance", 
        "P/U City", "Drop Address/Entrance", "Drop City", "Miles", 
        "Fare", "Trip Type", "StandingOrder Id", "One Way", "Comments"
    ]
    
    # Check if the DataFrame has enough columns
    if len(df.columns) < len(headers):
        # If not, extend the DataFrame with empty columns
        for i in range(len(df.columns), len(headers)):
            df[f"Column{i}"] = ""
    
    # Rename the first len(headers) columns
    df.columns = headers + list(df.columns[len(headers):])
    
    # Step 2: Fix data alignment for rows with "DAR" or "Taxi" in Trip Accepted
    fixed_rows = []
    
    for _, row in df.iterrows():
        row_data = row.tolist()
        if row_data[0] in ["DAR", "Taxi"]:
            # Insert blank cell at beginning and shift right
            fixed_rows.append([None] + row_data[:-1])
        else:
            fixed_rows.append(row_data)
    
    # Create new dataframe with fixed alignment
    df = pd.DataFrame(fixed_rows, columns=df.columns)
    
    # Step 2.5: Fix phone column misalignment
    df = fix_phone_column_misalignment(df)
    
    # Step 3: Add monitoring indicator to names based on comments
    for idx, row in df.iterrows():
        comments = str(row["Comments"]).upper()  # Convert to string and uppercase for case-insensitive comparison
        if "MONITOR" in comments or "MT" in comments:
            df.at[idx, "Name"] = f"{row['Name']} *monitor"
    
    # Step 4: Delete specified columns
    columns_to_keep = ["Name", "Phone", "P/U Time", "Appt Time", "P/U Address/Entrance", 
                    "P/U City", "Drop Address/Entrance", "Drop City"]
    df = df[columns_to_keep]
    
    # Step 5: Group identical names together and create group numbers
    unique_names = []
    group_numbers = []
    
    for name in df["Name"]:
        if name not in unique_names:
            unique_names.append(name)
            group_numbers.append(len(unique_names))
        else:
            # Add an empty string for repeated names
            group_numbers.append("")
    
    # Insert group numbers as the first column
    df.insert(0, "Group", group_numbers)
    
    # Create Word document
    doc = Document()
    
    # Set landscape orientation
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    # Add introduction text
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add title
    title_run = intro.add_run(f"RELIAMED {tomorrow_formatted}, FINAL SCHEDULED TRIPS\n")
    title_run.bold = True
    title_run.font.size = Pt(14)
    
    # Add note and dispatch info
    note_run = intro.add_run("*Schedule is subject to change\nDispatch 508-981-5919")
    note_run.font.size = Pt(11)
    
    # Add spacing after introduction
    doc.add_paragraph()
    
    # Create table in Word document
    table = doc.add_table(rows=len(df)+1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Add headers with bold formatting
    for i, column in enumerate(df.columns):
        cell = table.cell(0, i)
        cell.text = column
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)
    
    # Add data
    for i, row in enumerate(df.itertuples(index=False)):
        for j, value in enumerate(row):
            # Convert NaN or None to empty string
            if pd.isna(value):
                cell_text = ""
            else:
                cell_text = str(value)
            table.cell(i+1, j).text = cell_text
    
    # Apply formatting to table
    for row in table.rows:
        for cell in row.cells:
            # Add borders
            cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders {}><w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/><w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/></w:tcBorders>'.format(nsdecls('w'))))
            
            # Set font and size for all cells
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Aptos"
                    run.font.size = Pt(11)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
    
    # Adjust column widths
    for column in table.columns:
        column.width = Inches(0.9)
    
    # Save document
    doc.save(output_file)
    
    print(f"Report successfully generated and saved as: {output_file}")

def fix_phone_column_misalignment(df):
    # Create a copy of the DataFrame to avoid modifying the original during iteration
    fixed_rows = []
    
    for _, row in df.iterrows():
        row_data = row.tolist()
        phone_value = str(row_data[3])  # Phone column is at index 3
        
        # Check if the phone value matches a date pattern (MM/DD/YYYY or similar)
        if re.match(r'\d{1,2}/\d{1,2}/\d{4}', phone_value):
            # Insert empty cell at phone position and shift right
            fixed_row = row_data[:3] + [None] + row_data[3:-1]
            fixed_rows.append(fixed_row)
        else:
            fixed_rows.append(row_data)
    
    # Create new dataframe with fixed alignment
    return pd.DataFrame(fixed_rows, columns=df.columns)

if __name__ == "__main__":
    process_excel_report()